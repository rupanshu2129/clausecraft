import io
import os
import pathlib
from typing import Tuple

from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document

MAX_CHARS = 250_000

def _clip(s: str, limit: int = MAX_CHARS) -> str:
    return s[:limit] if len(s) > limit else s

def extract_text_from_pdf(data: bytes) -> str:
    with io.BytesIO(data) as fh:
        txt = pdf_extract_text(fh) or ""
    return _clip(txt)

def extract_text_from_docx(data: bytes) -> str:
    with io.BytesIO(data) as fh:
        doc = Document(fh)
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_txt = " | ".join((cell.text or "").strip() for cell in row.cells)
                if row_txt.strip():
                    parts.append(row_txt)
    return _clip("\n".join(parts))

def extract_text(buf: bytes, filename: str) -> Tuple[str, str]:
    ext = pathlib.Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(buf), "pdf"
    if ext == ".docx":
        return extract_text_from_docx(buf), "docx"
    if ext == ".txt":
        return _clip(buf.decode("utf-8", errors="ignore")), "txt"
    return _clip(buf.decode("utf-8", errors="ignore")), "unknown"
